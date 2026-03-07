from __future__ import annotations

from dataclasses import dataclass
from functools import lru_cache
from html import escape
from io import BytesIO
from pathlib import Path
import re


@dataclass(frozen=True)
class TextRun:
    text: str
    bold: bool = False
    italic: bool = False
    code: bool = False
    link: str | None = None


@dataclass(frozen=True)
class HeadingBlock:
    level: int
    runs: list[TextRun]


@dataclass(frozen=True)
class ParagraphBlock:
    runs: list[TextRun]


@dataclass(frozen=True)
class ListBlock:
    ordered: bool
    items: list[list[TextRun]]


@dataclass(frozen=True)
class QuoteBlock:
    blocks: list[Block]


@dataclass(frozen=True)
class CodeBlock:
    code: str
    language: str | None = None


@dataclass(frozen=True)
class TableBlock:
    headers: list[list[TextRun]]
    rows: list[list[list[TextRun]]]


@dataclass(frozen=True)
class HorizontalRuleBlock:
    pass


Block = (
    HeadingBlock
    | ParagraphBlock
    | ListBlock
    | QuoteBlock
    | CodeBlock
    | TableBlock
    | HorizontalRuleBlock
)


@dataclass(frozen=True)
class DocumentModel:
    blocks: list[Block]


@dataclass(frozen=True)
class PageTheme:
    margin_top: int = 48
    margin_right: int = 48
    margin_bottom: int = 48
    margin_left: int = 48


@dataclass(frozen=True)
class BodyTheme:
    font_family: str = "Helvetica"
    bold_font_family: str = "Helvetica-Bold"
    italic_font_family: str = "Helvetica-Oblique"
    bold_italic_font_family: str = "Helvetica-BoldOblique"
    docx_font_family: str = "Arial"
    docx_east_asia_font_family: str = "SimSun"
    font_size: int = 11
    line_height: int = 15
    color: str = "#111827"
    space_after: int = 10


@dataclass(frozen=True)
class HeadingTheme:
    font_size: int
    space_before: int
    space_after: int


@dataclass(frozen=True)
class InlineCodeTheme:
    font_family: str = "Courier"
    docx_font_family: str = "Courier New"
    docx_east_asia_font_family: str = "SimSun"
    background: str = "#F3F4F6"


@dataclass(frozen=True)
class CodeBlockTheme:
    font_family: str = "Courier"
    docx_font_family: str = "Courier New"
    docx_east_asia_font_family: str = "SimSun"
    font_size: int = 9
    line_height: int = 12
    background: str = "#F6F8FA"
    border_color: str = "#D0D7DE"
    padding: int = 8


@dataclass(frozen=True)
class TableTheme:
    header_background: str = "#EEF2F7"
    border_color: str = "#CBD5E1"
    cell_padding: int = 6


@dataclass(frozen=True)
class BlockquoteTheme:
    text_color: str = "#374151"
    left_indent: int = 18


@dataclass(frozen=True)
class ExportTheme:
    page: PageTheme
    body: BodyTheme
    headings: dict[int, HeadingTheme]
    inline_code: InlineCodeTheme
    code_block: CodeBlockTheme
    table: TableTheme
    blockquote: BlockquoteTheme


@dataclass(frozen=True)
class ExportedDocument:
    content: bytes
    media_type: str
    extension: str


FENCED_CODE_RE = re.compile(r"^(?P<fence>`{3,}|~{3,})(?P<info>[^`]*)$")
ATX_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
UNORDERED_LIST_RE = re.compile(r"^\s*[-+*]\s+(.*)$")
ORDERED_LIST_RE = re.compile(r"^\s*(\d+)\.\s+(.*)$")
IMAGE_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")
CODE_SPAN_RE = re.compile(r"`([^`]+)`")
STRONG_RE = re.compile(r"(\*\*.+?\*\*|__.+?__)")
INLINE_TOKEN_RE = re.compile(
    r"(!\[[^\]]*\]\([^)]+\)|\[[^\]]+\]\([^)]+\)|`[^`]+`|\*\*.+?\*\*|__.+?__|\*[^*\n]+\*|_[^_\n]+_)",
)
CJK_TEXT_RE = re.compile(r"[\u2E80-\u2EFF\u3000-\u303F\u3040-\u30FF\u3100-\u312F\u31A0-\u31BF\u3400-\u4DBF\u4E00-\u9FFF\uAC00-\uD7AF\uF900-\uFAFF\uFE30-\uFE4F\uFF00-\uFFEF]")
PDF_CJK_FONT_NAME = "STSong-Light"
PDF_PANEL_SANS_FONT_FILES = {
    "regular": (
        "Geist-Regular.ttf",
        "Geist.ttf",
        "Inter-Regular.ttf",
        "Arial.ttf",
        "LiberationSans-Regular.ttf",
        "DejaVuSans.ttf",
        "NotoSans-Regular.ttf",
    ),
    "bold": (
        "Geist-Bold.ttf",
        "Inter-Bold.ttf",
        "Arial Bold.ttf",
        "Arialbd.ttf",
        "LiberationSans-Bold.ttf",
        "DejaVuSans-Bold.ttf",
        "NotoSans-Bold.ttf",
    ),
    "italic": (
        "Geist-Regular.ttf",
        "Inter-Italic.ttf",
        "Arial Italic.ttf",
        "Ariali.ttf",
        "LiberationSans-Italic.ttf",
        "DejaVuSans-Oblique.ttf",
        "NotoSans-Italic.ttf",
    ),
    "bold_italic": (
        "Geist-Bold.ttf",
        "Inter-BoldItalic.ttf",
        "Arial Bold Italic.ttf",
        "Arialbi.ttf",
        "LiberationSans-BoldItalic.ttf",
        "DejaVuSans-BoldOblique.ttf",
        "NotoSans-BoldItalic.ttf",
    ),
}
PDF_FONT_SEARCH_ROOTS = (
    Path("/usr/share/fonts"),
    Path("/usr/local/share/fonts"),
    Path.home() / ".fonts",
    Path("/root/super-deer-flow"),
)


def get_default_export_theme() -> ExportTheme:
    return ExportTheme(
        page=PageTheme(),
        body=BodyTheme(),
        headings={
            1: HeadingTheme(font_size=22, space_before=10, space_after=12),
            2: HeadingTheme(font_size=18, space_before=8, space_after=10),
            3: HeadingTheme(font_size=16, space_before=6, space_after=8),
            4: HeadingTheme(font_size=14, space_before=4, space_after=6),
            5: HeadingTheme(font_size=12, space_before=4, space_after=6),
            6: HeadingTheme(font_size=11, space_before=4, space_after=6),
        },
        inline_code=InlineCodeTheme(),
        code_block=CodeBlockTheme(),
        table=TableTheme(),
        blockquote=BlockquoteTheme(),
    )


def export_markdown_document(markdown_text: str, format: str) -> ExportedDocument:
    document = parse_markdown_document(markdown_text)
    theme = get_default_export_theme()
    if format == "pdf":
        return ExportedDocument(
            content=render_pdf(document, theme),
            media_type="application/pdf",
            extension="pdf",
        )
    if format == "docx":
        return ExportedDocument(
            content=render_docx(document, theme),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            extension="docx",
        )
    raise ValueError(f"Unsupported export format: {format}")


def parse_markdown_document(markdown_text: str) -> DocumentModel:
    lines = markdown_text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    return DocumentModel(blocks=_parse_blocks(lines))


def _parse_blocks(lines: list[str]) -> list[Block]:
    blocks: list[Block] = []
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue

        fenced_match = FENCED_CODE_RE.match(line.strip())
        if fenced_match:
            index, block = _parse_fenced_code(lines, index, fenced_match)
            blocks.append(block)
            continue

        heading_match = ATX_HEADING_RE.match(line.strip())
        if heading_match:
            level = len(heading_match.group(1))
            blocks.append(HeadingBlock(level=level, runs=parse_inline_runs(heading_match.group(2).strip())))
            index += 1
            continue

        if _is_horizontal_rule(line):
            blocks.append(HorizontalRuleBlock())
            index += 1
            continue

        if _is_table_start(lines, index):
            index, block = _parse_table(lines, index)
            blocks.append(block)
            continue

        if line.lstrip().startswith(">"):
            index, block = _parse_blockquote(lines, index)
            blocks.append(block)
            continue

        if UNORDERED_LIST_RE.match(line) or ORDERED_LIST_RE.match(line):
            index, block = _parse_list(lines, index)
            blocks.append(block)
            continue

        index, block = _parse_paragraph(lines, index)
        blocks.append(block)

    return blocks


def _parse_fenced_code(lines: list[str], index: int, match: re.Match[str]) -> tuple[int, CodeBlock]:
    fence = match.group("fence")
    info = match.group("info").strip() or None
    code_lines: list[str] = []
    index += 1
    while index < len(lines):
        line = lines[index]
        if line.strip().startswith(fence):
            index += 1
            break
        code_lines.append(line)
        index += 1
    return index, CodeBlock(code="\n".join(code_lines), language=info)


def _parse_blockquote(lines: list[str], index: int) -> tuple[int, QuoteBlock]:
    quote_lines: list[str] = []
    while index < len(lines):
        line = lines[index]
        stripped = line.lstrip()
        if not stripped.startswith(">"):
            break
        quote_lines.append(re.sub(r"^\s*>\s?", "", line))
        index += 1
    return index, QuoteBlock(blocks=_parse_blocks(quote_lines))


def _parse_list(lines: list[str], index: int) -> tuple[int, ListBlock]:
    ordered = bool(ORDERED_LIST_RE.match(lines[index]))
    items: list[list[TextRun]] = []

    while index < len(lines):
        line = lines[index]
        match = ORDERED_LIST_RE.match(line) if ordered else UNORDERED_LIST_RE.match(line)
        if not match:
            break

        item_lines = [match.group(2 if ordered else 1).strip()]
        index += 1

        while index < len(lines):
            next_line = lines[index]
            if not next_line.strip():
                peek = index + 1
                while peek < len(lines) and not lines[peek].strip():
                    peek += 1
                if peek < len(lines):
                    next_match = ORDERED_LIST_RE.match(lines[peek]) if ordered else UNORDERED_LIST_RE.match(lines[peek])
                    if next_match:
                        index = peek
                        break
                index = peek
                break
            if (ORDERED_LIST_RE.match(next_line) if ordered else UNORDERED_LIST_RE.match(next_line)):
                break
            if _starts_new_block(next_line) and not next_line.startswith((" ", "\t")):
                break
            item_lines.append(next_line.strip())
            index += 1

        items.append(parse_inline_runs(" ".join(part for part in item_lines if part)))

    return index, ListBlock(ordered=ordered, items=items)


def _parse_paragraph(lines: list[str], index: int) -> tuple[int, ParagraphBlock]:
    paragraph_lines = [lines[index].strip()]
    index += 1
    while index < len(lines):
        line = lines[index]
        if not line.strip() or _starts_new_block(line):
            break
        paragraph_lines.append(line.strip())
        index += 1
    return index, ParagraphBlock(runs=parse_inline_runs(" ".join(paragraph_lines)))


def _parse_table(lines: list[str], index: int) -> tuple[int, TableBlock]:
    header_cells = _split_table_row(lines[index])
    index += 2
    rows: list[list[list[TextRun]]] = []
    while index < len(lines):
        line = lines[index]
        if not line.strip() or "|" not in line:
            break
        rows.append([parse_inline_runs(cell) for cell in _split_table_row(line)])
        index += 1

    column_count = max([len(header_cells), *(len(row) for row in rows)])
    header_cells = _pad_cells(header_cells, column_count)
    rows = [_pad_run_cells(row, column_count) for row in rows]

    return index, TableBlock(
        headers=[parse_inline_runs(cell) for cell in header_cells],
        rows=rows,
    )


def _pad_cells(cells: list[str], count: int) -> list[str]:
    return cells + ([""] * (count - len(cells)))


def _pad_run_cells(cells: list[list[TextRun]], count: int) -> list[list[TextRun]]:
    return cells + ([[]] * (count - len(cells)))


def _split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip() for cell in stripped.split("|")]


def _is_table_start(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    if "|" not in lines[index]:
        return False
    alignment_cells = _split_table_row(lines[index + 1])
    return bool(alignment_cells) and all(_is_table_alignment_cell(cell) for cell in alignment_cells)


def _is_table_alignment_cell(cell: str) -> bool:
    stripped = cell.strip()
    return bool(stripped) and re.fullmatch(r":?-{3,}:?", stripped) is not None


def _is_horizontal_rule(line: str) -> bool:
    stripped = line.strip().replace(" ", "")
    return stripped in {"***", "---", "___"} or re.fullmatch(r"([*_\-])\1{2,}", stripped) is not None


def _starts_new_block(line: str) -> bool:
    stripped = line.strip()
    if not stripped:
        return True
    return bool(
        FENCED_CODE_RE.match(stripped)
        or ATX_HEADING_RE.match(stripped)
        or UNORDERED_LIST_RE.match(line)
        or ORDERED_LIST_RE.match(line)
        or stripped.startswith(">")
        or _is_horizontal_rule(line)
    )


def parse_inline_runs(text: str) -> list[TextRun]:
    if not text:
        return []

    runs: list[TextRun] = []
    position = 0
    for match in INLINE_TOKEN_RE.finditer(text):
        if match.start() > position:
            runs.append(TextRun(text=text[position : match.start()]))

        token = match.group(0)
        runs.extend(_token_to_runs(token))
        position = match.end()

    if position < len(text):
        runs.append(TextRun(text=text[position:]))

    return _merge_runs(runs)


def _token_to_runs(token: str) -> list[TextRun]:
    image_match = IMAGE_RE.fullmatch(token)
    if image_match:
        alt_text = image_match.group(1).strip() or "Image"
        url = image_match.group(2).strip()
        return [TextRun(text=f"{alt_text} ({url})", link=url)]

    link_match = LINK_RE.fullmatch(token)
    if link_match:
        label = link_match.group(1)
        url = link_match.group(2).strip()
        return _apply_link(parse_inline_runs(label), url)

    code_match = CODE_SPAN_RE.fullmatch(token)
    if code_match:
        return [TextRun(text=code_match.group(1), code=True)]

    if STRONG_RE.fullmatch(token):
        return _apply_text_style(parse_inline_runs(token[2:-2]), bold=True)

    if token.startswith("*") and token.endswith("*") and len(token) >= 2:
        return _apply_text_style(parse_inline_runs(token[1:-1]), italic=True)

    if token.startswith("_") and token.endswith("_") and len(token) >= 2:
        return _apply_text_style(parse_inline_runs(token[1:-1]), italic=True)

    return [TextRun(text=token)]


def _apply_text_style(runs: list[TextRun], *, bold: bool = False, italic: bool = False) -> list[TextRun]:
    if not runs:
        return []
    return [
        TextRun(
            text=run.text,
            bold=run.bold or bold,
            italic=run.italic or italic,
            code=run.code,
            link=run.link,
        )
        for run in runs
    ]


def _apply_link(runs: list[TextRun], url: str) -> list[TextRun]:
    if not runs:
        return [TextRun(text=url, link=url)]
    return [
        TextRun(
            text=run.text,
            bold=run.bold,
            italic=run.italic,
            code=run.code,
            link=url,
        )
        for run in runs
    ]


def _merge_runs(runs: list[TextRun]) -> list[TextRun]:
    merged: list[TextRun] = []
    for run in runs:
        if not run.text:
            continue
        if merged and merged[-1].bold == run.bold and merged[-1].italic == run.italic and merged[-1].code == run.code and merged[-1].link == run.link:
            previous = merged[-1]
            merged[-1] = TextRun(
                text=previous.text + run.text,
                bold=previous.bold,
                italic=previous.italic,
                code=previous.code,
                link=previous.link,
            )
            continue
        merged.append(run)
    return merged


def render_pdf(document: DocumentModel, theme: ExportTheme) -> bytes:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import HRFlowable, Paragraph, Preformatted, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError("PDF export dependencies are not installed") from exc

    _ensure_pdf_cjk_font_registered()

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=theme.page.margin_top,
        rightMargin=theme.page.margin_right,
        bottomMargin=theme.page.margin_bottom,
        leftMargin=theme.page.margin_left,
    )
    styles = getSampleStyleSheet()
    body_style = ParagraphStyle(
        "ExportBody",
        parent=styles["BodyText"],
        fontName=_pdf_body_font_name(theme),
        fontSize=theme.body.font_size,
        leading=theme.body.line_height,
        textColor=_pdf_color(colors, theme.body.color),
        spaceAfter=theme.body.space_after,
    )
    heading_styles = {
        level: ParagraphStyle(
            f"ExportHeading{level}",
            parent=styles["Heading1"],
            fontName=_pdf_body_font_name(theme, bold=True),
            fontSize=heading_theme.font_size,
            leading=heading_theme.font_size + 4,
            spaceBefore=heading_theme.space_before,
            spaceAfter=heading_theme.space_after,
            textColor=_pdf_color(colors, theme.body.color),
        )
        for level, heading_theme in theme.headings.items()
    }
    code_style = ParagraphStyle(
        "ExportCode",
        parent=styles["Code"],
        fontName=theme.code_block.font_family,
        fontSize=theme.code_block.font_size,
        leading=theme.code_block.line_height,
        textColor=_pdf_color(colors, theme.body.color),
    )

    story = []
    for block in document.blocks:
        _append_pdf_block(
            story,
            block,
            theme,
            body_style=body_style,
            heading_styles=heading_styles,
            code_style=code_style,
            paragraph_style_cls=ParagraphStyle,
            colors_module=colors,
            paragraph_cls=Paragraph,
            preformatted_cls=Preformatted,
            table_cls=Table,
            table_style_cls=TableStyle,
            spacer_cls=Spacer,
            hr_cls=HRFlowable,
            mm=mm,
            quote_depth=0,
        )

    if not story:
        story.append(Paragraph("&nbsp;", body_style))

    doc.build(story)
    return buffer.getvalue()


def _append_pdf_block(
    story: list,
    block: Block,
    theme: ExportTheme,
    *,
    body_style,
    heading_styles,
    code_style,
    paragraph_style_cls,
    colors_module,
    paragraph_cls,
    preformatted_cls,
    table_cls,
    table_style_cls,
    spacer_cls,
    hr_cls,
    mm,
    quote_depth: int,
) -> None:
    quote_indent = quote_depth * theme.blockquote.left_indent
    if isinstance(block, HeadingBlock):
        style = heading_styles.get(block.level, heading_styles[6])
        if quote_indent:
            style = paragraph_style_cls(
                f"QuoteHeading{block.level}-{quote_depth}",
                parent=style,
                leftIndent=quote_indent,
                textColor=_pdf_color(colors_module, theme.blockquote.text_color),
            )
        story.append(paragraph_cls(_runs_to_reportlab_markup(block.runs, theme), style))
        return

    if isinstance(block, ParagraphBlock):
        style = body_style
        if quote_indent:
            style = paragraph_style_cls(
                f"QuoteBody-{quote_depth}",
                parent=body_style,
                leftIndent=quote_indent,
                textColor=_pdf_color(colors_module, theme.blockquote.text_color),
            )
        story.append(paragraph_cls(_runs_to_reportlab_markup(block.runs, theme), style))
        return

    if isinstance(block, ListBlock):
        for index, item in enumerate(block.items, start=1):
            prefix = f"{index}. " if block.ordered else "• "
            style = paragraph_style_cls(
                f"List-{quote_depth}-{index}",
                parent=body_style,
                leftIndent=quote_indent + 16,
                firstLineIndent=0,
            )
            if quote_indent:
                style.textColor = _pdf_color(colors_module, theme.blockquote.text_color)
            story.append(paragraph_cls(escape(prefix) + _runs_to_reportlab_markup(item, theme), style))
        story.append(spacer_cls(1, 2 * mm))
        return

    if isinstance(block, QuoteBlock):
        for nested in block.blocks:
            _append_pdf_block(
                story,
                nested,
                theme,
                body_style=body_style,
                heading_styles=heading_styles,
                code_style=code_style,
                paragraph_style_cls=paragraph_style_cls,
                colors_module=colors_module,
                paragraph_cls=paragraph_cls,
                preformatted_cls=preformatted_cls,
                table_cls=table_cls,
                table_style_cls=table_style_cls,
                spacer_cls=spacer_cls,
                hr_cls=hr_cls,
                mm=mm,
                quote_depth=quote_depth + 1,
            )
        return

    if isinstance(block, CodeBlock):
        block_code_style = code_style
        if _contains_cjk_text(block.code):
            cjk_font_name = _ensure_pdf_cjk_font_registered()
            if cjk_font_name:
                block_code_style = paragraph_style_cls(
                    f"ExportCodeCjk-{quote_depth}",
                    parent=code_style,
                    fontName=cjk_font_name,
                )
        pre = preformatted_cls(block.code or " ", block_code_style)
        code_table = table_cls([[pre]])
        code_table.setStyle(
            table_style_cls(
                [
                    ("BACKGROUND", (0, 0), (-1, -1), _pdf_color(colors_module, theme.code_block.background)),
                    ("BOX", (0, 0), (-1, -1), 0.75, _pdf_color(colors_module, theme.code_block.border_color)),
                    ("LEFTPADDING", (0, 0), (-1, -1), theme.code_block.padding),
                    ("RIGHTPADDING", (0, 0), (-1, -1), theme.code_block.padding),
                    ("TOPPADDING", (0, 0), (-1, -1), theme.code_block.padding),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), theme.code_block.padding),
                ]
            )
        )
        story.append(code_table)
        story.append(spacer_cls(1, 3 * mm))
        return

    if isinstance(block, TableBlock):
        rows = [block.headers, *block.rows]
        table_data = []
        for row in rows:
            table_data.append([
                paragraph_cls(_runs_to_reportlab_markup(cell, theme) or "&nbsp;", body_style)
                for cell in row
            ])

        table = table_cls(table_data, repeatRows=1)
        table.setStyle(
            table_style_cls(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), _pdf_color(colors_module, theme.table.header_background)),
                    ("GRID", (0, 0), (-1, -1), 0.5, _pdf_color(colors_module, theme.table.border_color)),
                    ("LEFTPADDING", (0, 0), (-1, -1), theme.table.cell_padding),
                    ("RIGHTPADDING", (0, 0), (-1, -1), theme.table.cell_padding),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)
        story.append(spacer_cls(1, 3 * mm))
        return

    if isinstance(block, HorizontalRuleBlock):
        story.append(hr_cls(width="100%", thickness=0.75, color=_pdf_color(colors_module, theme.table.border_color)))
        story.append(spacer_cls(1, 3 * mm))


def _pdf_color(colors_module, hex_value: str):
    return colors_module.HexColor(hex_value)


def _contains_cjk_text(text: str) -> bool:
    return bool(text and CJK_TEXT_RE.search(text))


def _find_first_matching_font_file(candidate_names: tuple[str, ...]) -> Path | None:
    normalized_candidates = {name.lower() for name in candidate_names}
    for root in PDF_FONT_SEARCH_ROOTS:
        if not root.exists():
            continue
        for font_path in root.rglob("*"):
            if not font_path.is_file():
                continue
            if font_path.suffix.lower() not in {".ttf", ".ttc", ".otf"}:
                continue
            if font_path.name.lower() in normalized_candidates:
                return font_path
    return None


def _discover_pdf_panel_font_paths() -> dict[str, Path | None]:
    return {
        variant: _find_first_matching_font_file(candidate_names)
        for variant, candidate_names in PDF_PANEL_SANS_FONT_FILES.items()
    }


@lru_cache(maxsize=1)
def _resolve_pdf_latin_font_names() -> dict[str, str]:
    fallback = {
        "regular": "Helvetica",
        "bold": "Helvetica-Bold",
        "italic": "Helvetica-Oblique",
        "bold_italic": "Helvetica-BoldOblique",
    }

    discovered = _discover_pdf_panel_font_paths()
    if not discovered.get("regular") or not discovered.get("bold"):
        return fallback

    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
    except ImportError:
        return fallback

    resolved: dict[str, str] = {}
    for variant, font_path in discovered.items():
        if not font_path:
            continue

        font_name = f"ExportPanelSans-{variant.replace('_', '-').title()}"
        try:
            pdfmetrics.getFont(font_name)
        except KeyError:
            pdfmetrics.registerFont(TTFont(font_name, str(font_path)))
        except Exception:
            return fallback
        resolved[variant] = font_name

    resolved.setdefault("regular", fallback["regular"])
    resolved.setdefault("bold", fallback["bold"])
    resolved.setdefault("italic", resolved["regular"])
    resolved.setdefault("bold_italic", resolved["bold"])
    return resolved


def _pdf_body_font_name(theme: ExportTheme, *, bold: bool = False, italic: bool = False) -> str:
    resolved = _resolve_pdf_latin_font_names()
    fallback = {
        "regular": theme.body.font_family,
        "bold": theme.body.bold_font_family,
        "italic": theme.body.italic_font_family,
        "bold_italic": theme.body.bold_italic_font_family,
    }

    if bold and italic:
        return resolved.get("bold_italic", fallback["bold_italic"])
    if bold:
        return resolved.get("bold", fallback["bold"])
    if italic:
        return resolved.get("italic", fallback["italic"])
    return resolved.get("regular", fallback["regular"])


@lru_cache(maxsize=1)
def _ensure_pdf_cjk_font_registered() -> str | None:
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    except ImportError:
        return None

    try:
        pdfmetrics.getFont(PDF_CJK_FONT_NAME)
    except KeyError:
        pdfmetrics.registerFont(UnicodeCIDFont(PDF_CJK_FONT_NAME))
    except Exception:
        return None

    return PDF_CJK_FONT_NAME


def _runs_to_reportlab_markup(runs: list[TextRun], theme: ExportTheme) -> str:
    fragments: list[str] = []
    cjk_font_name = _ensure_pdf_cjk_font_registered()

    for run in runs:
        plain_text = run.text
        if run.link and run.link not in run.text:
            plain_text = f"{plain_text} ({run.link})"

        text = escape(plain_text).replace("\n", "<br/>")
        if not text:
            continue

        if run.code:
            font_name = cjk_font_name if cjk_font_name and _contains_cjk_text(plain_text) else theme.inline_code.font_family
        elif cjk_font_name and _contains_cjk_text(plain_text):
            font_name = cjk_font_name
        else:
            font_name = _pdf_body_font_name(theme, bold=run.bold, italic=run.italic)

        text = f'<font name="{escape(font_name)}">{text}</font>'
        fragments.append(text)

    return "".join(fragments) or "&nbsp;"


def render_docx(document: DocumentModel, theme: ExportTheme) -> bytes:
    try:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError("Word export dependencies are not installed") from exc

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(theme.page.margin_top)
    section.right_margin = Pt(theme.page.margin_right)
    section.bottom_margin = Pt(theme.page.margin_bottom)
    section.left_margin = Pt(theme.page.margin_left)

    for block in document.blocks:
        _append_docx_block(
            doc,
            block,
            theme,
            quote_depth=0,
            table_alignment=WD_TABLE_ALIGNMENT,
            oxml_element=OxmlElement,
            qname=qn,
            pt=Pt,
            rgb_color=RGBColor,
        )

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _append_docx_block(
    doc,
    block: Block,
    theme: ExportTheme,
    *,
    quote_depth: int,
    table_alignment,
    oxml_element,
    qname,
    pt,
    rgb_color,
) -> None:
    if isinstance(block, HeadingBlock):
        paragraph = doc.add_paragraph(style=_docx_heading_style_name(block.level))
        _apply_docx_paragraph_base(paragraph, theme, pt, quote_depth, is_quote=quote_depth > 0)
        heading_theme = theme.headings.get(block.level, theme.headings[6])
        _append_docx_runs(paragraph, block.runs, theme, pt, rgb_color, oxml_element, qname, font_size=heading_theme.font_size)
        paragraph.paragraph_format.space_before = pt(heading_theme.space_before)
        paragraph.paragraph_format.space_after = pt(heading_theme.space_after)
        return

    if isinstance(block, ParagraphBlock):
        paragraph = doc.add_paragraph()
        _apply_docx_paragraph_base(paragraph, theme, pt, quote_depth, is_quote=quote_depth > 0)
        _append_docx_runs(paragraph, block.runs, theme, pt, rgb_color, oxml_element, qname)
        return

    if isinstance(block, ListBlock):
        list_style = _docx_list_style_name(block.ordered)
        for item in block.items:
            paragraph = doc.add_paragraph(style=list_style)
            _apply_docx_paragraph_base(paragraph, theme, pt, quote_depth, is_quote=quote_depth > 0)
            _append_docx_runs(paragraph, item, theme, pt, rgb_color, oxml_element, qname)
        return

    if isinstance(block, QuoteBlock):
        for nested in block.blocks:
            _append_docx_block(
                doc,
                nested,
                theme,
                quote_depth=quote_depth + 1,
                table_alignment=table_alignment,
                oxml_element=oxml_element,
                qname=qname,
                pt=pt,
                rgb_color=rgb_color,
            )
        return

    if isinstance(block, CodeBlock):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = table_alignment.LEFT
        cell = table.cell(0, 0)
        _set_docx_cell_background(cell, theme.code_block.background, oxml_element, qname)
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = pt(theme.body.space_after)
        run = paragraph.add_run(block.code or " ")
        _set_docx_run_fonts(
            run,
            ascii_font_family=theme.code_block.docx_font_family,
            east_asia_font_family=theme.code_block.docx_east_asia_font_family,
            oxml_element=oxml_element,
            qname=qname,
        )
        run.font.size = pt(theme.code_block.font_size)
        cell_paragraph = cell.paragraphs[0]
        cell_paragraph.paragraph_format.left_indent = pt(0)
        return

    if isinstance(block, TableBlock):
        table = doc.add_table(rows=len(block.rows) + 1, cols=len(block.headers), style="Table Grid")
        table.alignment = table_alignment.LEFT
        for col_index, cell_runs in enumerate(block.headers):
            cell = table.rows[0].cells[col_index]
            _set_docx_cell_background(cell, theme.table.header_background, oxml_element, qname)
            paragraph = cell.paragraphs[0]
            _append_docx_runs(paragraph, cell_runs, theme, pt, rgb_color, oxml_element, qname, font_size=theme.body.font_size)
        for row_index, row in enumerate(block.rows, start=1):
            for col_index, cell_runs in enumerate(row):
                paragraph = table.rows[row_index].cells[col_index].paragraphs[0]
                _append_docx_runs(paragraph, cell_runs, theme, pt, rgb_color, oxml_element, qname, font_size=theme.body.font_size)
        doc.add_paragraph()
        return

    if isinstance(block, HorizontalRuleBlock):
        paragraph = doc.add_paragraph()
        run = paragraph.add_run("―" * 24)
        run.font.color.rgb = rgb_color.from_string(theme.table.border_color.lstrip("#"))
        paragraph.paragraph_format.space_after = pt(theme.body.space_after)


def _append_docx_runs(paragraph, runs: list[TextRun], theme: ExportTheme, pt, rgb_color, oxml_element, qname, font_size: int | None = None) -> None:
    default_size = font_size or theme.body.font_size
    body_color = rgb_color.from_string(theme.body.color.lstrip("#"))
    quote_color = rgb_color.from_string(theme.blockquote.text_color.lstrip("#"))

    for run_data in runs or [TextRun(text="")]:
        text = run_data.text
        if run_data.link and run_data.link not in run_data.text:
            text = f"{text} ({run_data.link})"
        run = paragraph.add_run(text)
        run.bold = run_data.bold
        run.italic = run_data.italic
        _set_docx_run_fonts(
            run,
            ascii_font_family=theme.inline_code.docx_font_family if run_data.code else theme.body.docx_font_family,
            east_asia_font_family=theme.inline_code.docx_east_asia_font_family if run_data.code else theme.body.docx_east_asia_font_family,
            oxml_element=oxml_element,
            qname=qname,
        )
        run.font.size = pt(default_size)
        run.font.color.rgb = quote_color if paragraph.paragraph_format.left_indent else body_color


def _apply_docx_paragraph_base(paragraph, theme: ExportTheme, pt, quote_depth: int, *, is_quote: bool) -> None:
    paragraph.paragraph_format.space_after = pt(theme.body.space_after)
    paragraph.paragraph_format.line_spacing = 1.2
    if is_quote:
        paragraph.paragraph_format.left_indent = pt(theme.blockquote.left_indent * quote_depth)


def _set_docx_run_fonts(run, *, ascii_font_family: str, east_asia_font_family: str, oxml_element, qname) -> None:
    run.font.name = ascii_font_family
    run_properties = run._element.get_or_add_rPr()
    fonts = run_properties.rFonts
    if fonts is None:
        fonts = oxml_element("w:rFonts")
        run_properties.append(fonts)
    fonts.set(qname("w:ascii"), ascii_font_family)
    fonts.set(qname("w:hAnsi"), ascii_font_family)
    fonts.set(qname("w:eastAsia"), east_asia_font_family)
    fonts.set(qname("w:cs"), east_asia_font_family)


def _set_docx_cell_background(cell, color: str, oxml_element, qname) -> None:
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = oxml_element("w:shd")
    shading.set(qname("w:fill"), color.lstrip("#"))
    cell_properties.append(shading)


def _docx_heading_style_name(level: int) -> str:
    return f"Heading {min(max(level, 1), 6)}"


def _docx_list_style_name(ordered: bool) -> str:
    return "List Number" if ordered else "List Bullet"
